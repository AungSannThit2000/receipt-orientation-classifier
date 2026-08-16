from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_PATH = PROJECT_ROOT / "docs" / "Receipt_Orientation_Classifier_Project_Report.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
DEEP_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BLACK = "17212B"
GOLD = "7A5A00"
GREEN = "1F5E52"
RED = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_SIDE_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_SIDE_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("end", CELL_SIDE_DXA),
    ):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, "Calibri", 9, MUTED)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.extend((field_begin, instruction, field_end))


def set_run_font(run, name: str, size: float, color: str = BLACK, bold=False, italic=False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = {
        "Title": (30, DARK_BLUE, 0, 8),
        "Subtitle": (14, MUTED, 0, 22),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DEEP_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in style_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name.startswith("Heading") or style_name == "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Code Block" not in doc.styles:
        code_style = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = doc.styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(BLACK)
    code_style.paragraph_format.left_indent = Inches(0.18)
    code_style.paragraph_format.right_indent = Inches(0.18)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(7)
    code_style.paragraph_format.line_spacing = 1.0

    if "Caption Project" not in doc.styles:
        caption = doc.styles.add_style("Caption Project", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = doc.styles["Caption Project"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False

    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Calibri"
    bullet_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet_style.font.size = Pt(11)
    bullet_style.font.color.rgb = RGBColor.from_string(BLACK)
    bullet_style.paragraph_format.left_indent = Inches(0.375)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.188)
    bullet_style.paragraph_format.space_after = Pt(4)
    bullet_style.paragraph_format.line_spacing = 1.25


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header.paragraphs[0]
    header.text = "Receipt Orientation Classifier"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], "Calibri", 8.5, MUTED, bold=True)
    right = header.add_run("   |   Implementation Report")
    set_run_font(right, "Calibri", 8.5, MUTED)

    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.clear()
    footer.paragraph_format.space_before = Pt(0)
    add_page_number(footer)


def add_numbering_definition(doc: Document, kind: str) -> int:
    del doc, kind
    return -1


def add_list_item(doc: Document, text: str, num_id: int, bold_prefix: str | None = None):
    del num_id
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    marker = paragraph.add_run("\u2022 ")
    set_run_font(marker, "Calibri", 11, BLACK)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_numbered_list(doc: Document, items: tuple[str, ...]) -> None:
    for index, text in enumerate(items, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.44)
        paragraph.paragraph_format.first_line_indent = Inches(-0.27)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        marker = paragraph.add_run(f"{index}. ")
        set_run_font(marker, "Calibri", 11, BLACK)
        paragraph.add_run(text)


def add_callout(doc: Document, label: str, text: str, color: str = DEEP_BLUE) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.16)
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), CALLOUT)
    p_pr.append(shading)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color)
    border.append(left)
    p_pr.append(border)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, "Calibri", 11, color, bold=True)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, "Calibri", 11, BLACK)


def add_code(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shading)
    paragraph.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, "Calibri", 9.5, DARK_BLUE, bold=True)
    for row_data in rows:
        row = table.add_row()
        cells = row.cells
        for index, value in enumerate(row_data):
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            set_run_font(run, "Calibri", 9.5, BLACK)
            if index == 0:
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc: Document, path: Path, caption: str, width: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    picture = run.add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", path.stem.replace("_", " ").title())
    caption_paragraph = doc.add_paragraph(caption, style="Caption Project")
    caption_paragraph.paragraph_format.keep_together = True


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("AI CONCEPTS | COMPUTER VISION PROJECT")
    set_run_font(run, "Calibri", 10.5, GOLD, bold=True)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Receipt Orientation\nClassifier")

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Object detection, transfer learning, and OCR verification\n"
        "for four-way receipt direction"
    )

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    rule.paragraph_format.left_indent = Inches(1.2)
    rule.paragraph_format.right_indent = Inches(1.2)
    rule.paragraph_format.space_after = Pt(28)

    metadata = (
        ("Prepared by", "Aung Sann Thit"),
        ("Course", "CSX4201 AI Concepts"),
        ("Institution", "Assumption University"),
        ("Project version", "Draft 3 - deployable hybrid system"),
        ("Date", "August 2026"),
    )
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, "Calibri", 10.5, MUTED, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, "Calibri", 10.5, BLACK)

    doc.add_page_break()


def add_section_roadmap(doc: Document, bullet_id: int) -> None:
    doc.add_heading("Document roadmap", level=1)
    add_callout(
        doc,
        "Purpose",
        "Explain the complete project from label design and data generation through training, hybrid inference, testing, GitHub packaging, and Streamlit deployment.",
    )
    sections = (
        "Executive summary and project objectives",
        "Problem definition and four orientation labels",
        "Dataset cleaning, canonical orientation, and leakage-safe splitting",
        "Receipt detection and preprocessing",
        "Model experiments, selection, and evaluation",
        "OCR verification and final decision policy",
        "Streamlit application and real-time prediction flow",
        "Source-code execution guide",
        "Results, limitations, and next work",
        "GitHub and Streamlit Community Cloud deployment",
    )
    for item in sections:
        add_list_item(doc, item, bullet_id)

    doc.add_heading("Project at a glance", level=2)
    add_table(
        doc,
        ["Item", "Implemented choice"],
        [
            ["Task", "Four-class quarter-turn receipt orientation"],
            ["Selected classifier", "Partially fine-tuned MobileNetV3 Small"],
            ["Receipt localization", "YOLO-World proposal plus OpenCV refinement"],
            ["Polarity verification", "Thai and English EasyOCR"],
            ["Deployment", "Streamlit application, local and Community Cloud"],
            ["Dataset publication", "Excluded; only aggregate reports and selected weights are public"],
        ],
        [2160, 7200],
    )


def add_executive_summary(doc: Document, bullet_id: int) -> None:
    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "The project detects how a receipt is oriented in a photograph and assigns one of four labels: upright, tilted right, upside down, or tilted left. The initial classifier performed well on source-group-held-out generated samples but confused opposite directions when receipt shape was symmetric or printed text became small. Draft 3 resolves that weakness with a hybrid pipeline: a visual model chooses the orientation axis and OCR determines which side is the top."
    )
    doc.add_paragraph(
        "The final application first normalizes EXIF metadata, locates the receipt, refines its paper boundary, creates the same 224 x 224 input used in training, runs MobileNetV3, and compares two opposite OCR directions. Strong OCR evidence may confirm or override the model. Weak evidence returns an explicit uncertain result and prevents automatic rotation."
    )
    add_callout(
        doc,
        "Main result",
        "The hybrid system reached 98.28% on the 348-sample generated test split, compared with 92.82% for the model-only production path. Two upside-down Thai views of one external receipt were also corrected by OCR after the CNN predicted upright.",
        GREEN,
    )
    doc.add_heading("Objectives", level=2)
    for text in (
        "Create balanced, consistently labeled quarter-turn classes from unlabeled upright receipt photographs.",
        "Prevent leakage by splitting on physical source receipt before generating rotations.",
        "Compare a small CNN with transfer learning and preserve every trained checkpoint and metric artifact.",
        "Improve camera-photo inference by detecting and cropping the receipt before classification.",
        "Use OCR only after classification so the target orientation is not normalized away.",
        "Expose evidence, uncertainty, and correction output in a usable Streamlit application.",
        "Publish reproducible code and selected weights without publishing the receipt dataset.",
    ):
        add_list_item(doc, text, bullet_id)


def add_problem_and_labels(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("2. Problem definition and labels", level=1)
    doc.add_paragraph(
        "This is a coarse orientation classification problem rather than exact deskew estimation. The observable target is the direction of the receipt's top edge relative to the camera. A complete 360-degree circle is divided into four 90-degree regions."
    )
    add_table(
        doc,
        ["Class", "Clockwise range", "Operational meaning"],
        [
            ["upright", "315-360 or 0-45 degrees", "Text reads normally from top to bottom"],
            ["tilted_right", "45-135 degrees", "Receipt top points toward the right"],
            ["upside_down", "135-225 degrees", "Receipt is inverted"],
            ["tilted_left", "225-315 degrees", "Receipt top points toward the left"],
        ],
        [1800, 2700, 4860],
    )
    add_callout(
        doc,
        "Boundary rule",
        "Generated samples use class-centered angles with controlled jitter and avoid exact 45-degree boundaries, where a human label would also be ambiguous.",
    )
    doc.add_heading("Why opposite directions are difficult", level=2)
    doc.add_paragraph(
        "After a narrow receipt is resized onto a square canvas, its silhouette provides strong vertical-versus-horizontal evidence but much weaker top-versus-bottom evidence. Logos, totals, and text direction carry the polarity signal. If those details are faint, blurred, unfamiliar, or too small, upright and upside down can look nearly identical to the classifier. The same effect appears between left and right."
    )


def add_data_pipeline(doc: Document, bullet_id: int) -> None:
    doc.add_heading("3. Dataset preparation and labeling", level=1)
    doc.add_paragraph(
        "The raw dataset contained upright receipt photographs without orientation labels. The preparation pipeline transformed those sources into a balanced classification dataset while keeping each physical receipt in exactly one split."
    )
    doc.add_heading("Source audit", level=2)
    add_numbered_list(doc, (
        "Scan supported image formats and reject unreadable files.",
        "Compute SHA-256 for exact duplicate detection.",
        "Remove two duplicates from 200 files, leaving 198 unique source receipts.",
        "Assign source groups to train, validation, and test before any rotations are generated.",
    ))
    add_table(
        doc,
        ["Split", "Source receipts", "Generated images", "Per class"],
        [
            ["Train", "139", "1,668", "417"],
            ["Validation", "30", "360", "90"],
            ["Test", "29", "348", "87"],
            ["Total", "198", "2,376", "594"],
        ],
        [2100, 2280, 2640, 2340],
    )
    doc.add_heading("Canonical top-side verification", level=2)
    doc.add_paragraph(
        "Before generating labels, OCR compares possible text directions to verify which edge is the top. Low-score or closely tied cases are blocked for manual review. Auditable corrections are stored in config/orientation_overrides.csv. OCR is appropriate here because it verifies source labels; it is not applied to normalize an uploaded test image before classification."
    )
    doc.add_heading("Balanced rotation generation", level=2)
    doc.add_paragraph(
        "Every approved source produces three images in each of the four classes. Rotation uses class-centered random jitter, preserves aspect ratio, and fills exposed regions with the same neutral gray used by inference. This creates 12 samples per source and exact class balance."
    )
    for text in (
        "Keep source_group unchanged for all derived images.",
        "Record generated filename, source filename, split, class, and angle in the manifest.",
        "Create QA grids for crops, orientations, fallbacks, and review-required sources.",
        "Never place rotations of one source receipt in multiple splits.",
    ):
        add_list_item(doc, text, bullet_id)


def add_preprocessing(doc: Document, bullet_id: int) -> None:
    doc.add_heading("4. Receipt detection and preprocessing", level=1)
    doc.add_paragraph(
        "Training and inference share src/receipt_preprocessing.py. This is important because a classifier trained on tightly cropped receipts can fail if production uploads contain hands, tables, or patterned backgrounds."
    )
    doc.add_heading("OpenCV extraction", level=2)
    for text in (
        "Apply EXIF transpose and convert to RGB.",
        "Build Otsu brightness, low-saturation, neutral-bright, and edge masks.",
        "Find candidate contours and approximate quadrilaterals.",
        "Score area, rectangularity, convexity, aspect ratio, centrality, and border contact.",
        "Perspective-correct the best valid paper boundary.",
        "Use a full-frame fallback when no contour is trustworthy.",
    ):
        add_list_item(doc, text, bullet_id)
    doc.add_heading("Real-time object detection", level=2)
    doc.add_paragraph(
        "Draft 3 adds a persisted YOLO-World checkpoint for receipt and document proposals. The highest suitable box is padded, then the OpenCV extractor runs inside that smaller proposal. The detector's fixed vocabulary is embedded in the checkpoint, so the large CLIP text encoder is not needed at application inference."
    )
    add_callout(
        doc,
        "Design intent",
        "Object detection solves localization, not top-side recognition. The CNN and OCR remain responsible for orientation after the crop is created.",
    )
    doc.add_heading("Model canvas", level=2)
    doc.add_paragraph(
        "The extracted receipt is fitted onto a 224 x 224 gray canvas without stretching. The canvas is normalized with the ImageNet mean and standard deviation stored in the selected checkpoint. No quarter-turn correction is applied before the CNN sees the image."
    )


def add_training(doc: Document, bullet_id: int) -> None:
    doc.add_heading("5. Model training and selection", level=1)
    doc.add_paragraph(
        "Two architectures and two crop policies were compared. The strict policy excluded full-frame fallback sources; the full policy retained all approved sources. Every experiment used the same grouped evaluation sets."
    )
    add_table(
        doc,
        ["Experiment", "Policy", "Best val", "Full test", "Macro F1", "Size"],
        [
            ["Simple CNN", "Full", "79.72%", "87.36%", "87.31%", "1.22 MB"],
            ["Simple CNN", "Strict", "83.67%", "88.79%", "88.78%", "1.22 MB"],
            ["MobileNetV3 Small", "Full", "93.33%", "92.24%", "92.26%", "5.94 MB"],
            ["MobileNetV3 Small", "Strict", "92.33%", "88.79%", "88.85%", "5.94 MB"],
        ],
        [2280, 1200, 1320, 1440, 1440, 1680],
    )
    add_callout(
        doc,
        "Selected model",
        "mobilenet_v3_small_finetune_full_best.pt. Transfer learning outperformed the CNN trained from scratch, and retaining fallback crops added useful appearance variation.",
        GREEN,
    )
    doc.add_heading("Training configuration", level=2)
    for text in (
        "Seed 4201 and deterministic PyTorch behavior.",
        "Cross-entropy loss with 0.05 label smoothing.",
        "AdamW optimizer with cosine learning-rate scheduling.",
        "ImageNet initialization for MobileNetV3 Small.",
        "Fine-tune the final three feature blocks and the classifier head.",
        "Early stopping based on validation performance.",
        "Store class order, image size, normalization, architecture, and weights in every checkpoint.",
    ):
        add_list_item(doc, text, bullet_id)
    add_figure(
        doc,
        PROJECT_ROOT / "reports" / "training" / "curves_mobilenet_v3_small_finetune_full.png",
        "Figure 1. Training and validation curves for the selected MobileNetV3 experiment.",
        5.9,
    )


def add_ocr(doc: Document, bullet_id: int) -> None:
    doc.add_heading("6. OCR verification and final decision", level=1)
    doc.add_paragraph(
        "The model's most important residual errors were within opposite pairs. EasyOCR provides an independent text-direction signal after classification, operating on the high-resolution extracted receipt rather than the 224-pixel tensor input."
    )
    doc.add_heading("Pairwise design", level=2)
    add_table(
        doc,
        ["Model-selected axis", "Candidate rotations", "Candidate labels"],
        [
            ["Vertical", "0 and 180 degrees", "Upright, upside down"],
            ["Horizontal", "90 and 270 degrees", "Tilted right, tilted left"],
        ],
        [2400, 2880, 4080],
    )
    doc.add_paragraph(
        "The verifier detects text once, keeps the 24 strongest text regions, transforms those boxes for the opposite direction, and recognizes Thai and English in CPU batches. Each direction score combines OCR confidence, readable length, alphabetic content, and receipt keyword matches."
    )
    add_code(
        doc,
        "direction_score = sum(confidence * min(text_length, 30)^0.85)\n"
        "                + alphabetic_bonus\n"
        "                + 7.5 * receipt_keyword_hits\n\n"
        "ocr_margin = (best_score - second_score) / max(best_score, 1)"
    )
    doc.add_page_break()
    doc.add_heading("Runtime decision rules", level=2)
    add_table(
        doc,
        ["Rule", "Required evidence", "Outcome"],
        [
            ["Strict OCR", "Score >= 5.0 and margin >= 0.35", "Confirm or override model"],
            ["Model + OCR consensus", "Same label; OCR score >= 4.0; OCR margin >= 0.50; model confidence >= 0.55; model margin >= 0.15", "Accept shared label"],
            ["High-confidence consensus", "Same label; OCR score >= 2.0; OCR margin >= 0.50; model confidence >= 0.95; model margin >= 0.80", "Accept shared label"],
            ["Inconclusive", "Neither rule passes", "Return Uncertain"],
        ],
        [1920, 4800, 2640],
    )
    for text in (
        "A weak OCR result cannot override a disagreeing model.",
        "A low absolute OCR score may confirm an exceptionally strong matching model result only when the OCR pair remains well separated.",
        "Uncertain is a deliberate abstention state, not a fifth orientation class.",
        "The current thresholds are provisional until a larger grouped real-photo validation set is available.",
    ):
        add_list_item(doc, text, bullet_id)


def add_app(doc: Document) -> None:
    doc.add_heading("7. Streamlit application and prediction flow", level=1)
    doc.add_paragraph(
        "app.py implements the complete user workflow. Resource-heavy objects are cached with st.cache_resource, and upload-based session keys prevent a tab change from repeating detector, classifier, or OCR work."
    )
    steps = (
        "Validate file type, byte content, pixel count, and EXIF orientation.",
        "Run receipt object detection; fall back to full-image contour extraction when needed.",
        "Create the classifier input and run four-class softmax inference.",
        "Choose the vertical or horizontal OCR pair from the model output.",
        "Confirm, override, accept consensus, or abstain according to the configured thresholds.",
        "Display final label, decision source, timings, crop evidence, probabilities, and OCR evidence.",
        "Rotate only a resolved receipt and provide a PNG correction download.",
    )
    add_numbered_list(doc, steps)
    add_figure(
        doc,
        PROJECT_ROOT / "docs" / "assets" / "app_prediction.png",
        "Figure 2. Verified local Streamlit result. OCR corrects an upside-down generated sample after a low-confidence upright model prediction.",
        6.3,
    )
    doc.add_heading("User-facing evidence", level=2)
    doc.add_paragraph(
        "The interface exposes the uploaded image, detection overlay, extracted receipt, model input, all class probabilities, OCR scores for both candidate directions, decision margin, recognized text preview, corrected receipt, and separate processing times. This allows the user to diagnose whether a wrong answer came from localization, visual classification, or text recognition."
    )


def add_code_guide(doc: Document, bullet_id: int) -> None:
    doc.add_page_break()
    doc.add_heading("8. Source-code execution guide", level=1)
    doc.add_paragraph(
        "The following inventory explains why each file exists and when it runs. Training utilities are retained for reproducibility even though Community Cloud uses only the runtime path."
    )
    rows = [
        ["app.py", "Runtime", "Streamlit UI, caching, orchestration, evidence, and download"],
        ["receipt_preprocessing.py", "Shared", "Contour masks, extraction, perspective crop, canvas, rotation"],
        ["realtime_receipt_detection.py", "Runtime", "YOLO proposal, OpenCV refinement, fallback, overlay"],
        ["inference.py", "Runtime", "Checkpoint validation, model rebuild, softmax, correction"],
        ["ocr_orientation.py", "Runtime/research", "EasyOCR preparation, region reuse, scoring, deskew evidence"],
        ["hybrid_orientation.py", "Runtime", "Candidate pairs, thresholds, override, consensus, abstention"],
        ["prepare_dataset.py", "Research", "Deduplication, grouped split, canonical checks, generation, QA"],
        ["build_384_dataset.py", "Research", "Higher-resolution rebuild and camera augmentation"],
        ["prepare_real_photos.py", "Research", "Grouped real-photo intake and detector normalization"],
        ["train_models.py", "Research", "Four experiments, early stopping, checkpoints, metrics, plots"],
        ["train_vertical_direction.py", "Optional research", "Two-class upright versus upside-down experiment"],
        ["evaluate_hybrid.py", "Research", "Validation threshold search and locked test evaluation"],
        ["evaluate_external_photos.py", "QA", "Production-path real-photo regression check"],
        ["preview_fallback_recrops.py", "QA", "Visual comparison of fallback extraction candidates"],
    ]
    add_table(doc, ["File", "Role", "Intention"], rows, [2700, 1740, 4920])

    doc.add_heading("Clean environment and app", level=2)
    add_code(
        doc,
        "python -m venv .venv\n"
        ".\\.venv\\Scripts\\Activate.ps1\n"
        "python -m pip install --upgrade pip\n"
        "python -m pip install -r requirements.txt\n"
        "python -m unittest discover -s tests -v\n"
        "python -m streamlit run app.py"
    )
    doc.add_heading("Full research sequence", level=2)
    add_code(
        doc,
        "python -m pip install -r requirements-training.txt\n"
        "python src/prepare_dataset.py --source-dir <receipt-image-folder> --overwrite\n"
        "python src/train_models.py --experiments simple_cnn_full simple_cnn_strict "
        "mobilenet_v3_small_finetune_full mobilenet_v3_small_finetune_strict\n"
        "python src/evaluate_hybrid.py\n"
        "python src/evaluate_external_photos.py"
    )
    add_callout(
        doc,
        "Execution boundary",
        "The public repository can run inference and tests with included checkpoints. Reproducing training requires a separately supplied receipt dataset because raw and generated receipt images are intentionally excluded.",
    )
    doc.add_heading("Configuration files", level=2)
    for text in (
        "config/hybrid_ocr_config.json controls languages, OCR size, batching, thresholds, and abstention.",
        "config/draft3_training.json records the prepared 384-pixel experiment and real-photo collection policy.",
        "config/orientation_overrides.csv stores manual canonical-orientation corrections.",
        ".streamlit/config.toml defines theme, upload size, headless behavior, and telemetry settings.",
    ):
        add_list_item(doc, text, bullet_id, bold_prefix=text.split(" controls")[0] if " controls" in text else None)


def add_results(doc: Document, bullet_id: int) -> None:
    doc.add_heading("9. Results and interpretation", level=1)
    add_table(
        doc,
        ["Evaluation", "Samples", "Model only", "Pairwise OCR", "Hybrid"],
        [
            ["Validation", "360", "92.50%", "96.94%", "97.22%"],
            ["Test", "348", "92.82%", "96.26%", "98.28%"],
        ],
        [2400, 1320, 1800, 1920, 1920],
    )
    doc.add_paragraph(
        "On the generated test split, the hybrid policy produced 22 OCR overrides. Nineteen repaired model errors, and none changed a correct prediction to an incorrect one under the validation-selected evaluation policy. Nineteen OCR comparisons were inconclusive and retained the model in that historical evaluation."
    )
    add_figure(
        doc,
        PROJECT_ROOT / "reports" / "hybrid" / "confusion_hybrid_test.png",
        "Figure 3. Hybrid confusion matrix on the source-group-held-out generated test set.",
        4.1,
    )
    doc.add_heading("Public generated samples", level=2)
    add_table(
        doc,
        ["Actual", "Base model", "Final hybrid", "Decision"],
        [
            ["Tilted left", "Tilted right, 44.2%", "Tilted left", "OCR override"],
            ["Tilted right", "Tilted left, 41.8%", "Tilted right", "OCR override"],
            ["Upright", "Upright, 48.8%", "Upright", "OCR confirmed"],
            ["Upside down", "Upright, 52.4%", "Upside down", "OCR override"],
        ],
        [2040, 2760, 2280, 2280],
    )
    doc.add_heading("External real-photo regression", level=2)
    doc.add_paragraph(
        "Two upside-down Thai photographs of one physical receipt were held outside training. Object detection produced useful crops, but the base CNN predicted upright in both views. Thai/English OCR selected upside down and overrode both predictions. This confirms the known failure was repaired for that receipt; it is not a real-world accuracy estimate because there is only one unique physical receipt group."
    )
    doc.add_heading("Limitations", level=2)
    for text in (
        "Training data is dominated by generated rotations of upright English receipts.",
        "Pairwise OCR cannot repair a model-selected wrong vertical/horizontal axis.",
        "Thai/English OCR is CPU intensive and adds approximately 15-28 seconds in measured examples.",
        "Cloud cold start is slower because EasyOCR weights must download and initialize.",
        "Runtime thresholds are provisional and require grouped real-photo calibration.",
        "The output is a quadrant, not an exact continuous rotation angle.",
    ):
        add_list_item(doc, text, bullet_id)


def add_testing_and_deployment(doc: Document, bullet_id: int) -> None:
    doc.add_heading("10. Verification and deployment", level=1)
    doc.add_heading("Automated verification", level=2)
    add_table(
        doc,
        ["Test module", "Coverage"],
        [
            ["test_inference.py", "Checkpoint contract, probabilities, model input, public demos"],
            ["test_hybrid_orientation.py", "Pairs, override, confirmation, consensus, fallback, abstention"],
            ["test_ocr_orientation.py", "Region limits, enhancement path, multilingual scoring"],
            ["test_realtime_receipt_detection.py", "Box selection, coordinate mapping, proposal fallback"],
        ],
        [3600, 5760],
    )
    add_callout(doc, "Verified result", "All 20 unit tests pass, Python compilation passes, and the local app completed both the public upside-down OCR override and a high-confidence tilted-left consensus regression.", GREEN)
    doc.add_heading("GitHub publication boundary", level=2)
    add_table(
        doc,
        ["Published", "Excluded"],
        [
            ["Source code and configuration", "Raw and generated receipt datasets"],
            ["Selected classifier and detector weights", "Downloaded EasyOCR model cache"],
            ["Aggregate metrics, histories, and confusion matrices", "Per-image OCR text prediction tables"],
            ["Generated fictional demo receipts", "Private real-photo diagnostics"],
            ["Markdown guides and this DOCX report", "Experimental checkpoints not used by the app"],
        ],
        [4680, 4680],
    )
    doc.add_heading("Streamlit Community Cloud steps", level=2)
    add_numbered_list(doc, (
        "Connect the GitHub account that administers the repository.",
        "Create an app from the main branch with app.py as the entrypoint.",
        "Choose Python 3.10 in Advanced settings.",
        "Deploy without secrets and monitor the dependency build logs.",
        "Open the public app and run a generated demo prediction.",
        "Expect the first OCR request to download Thai, English, and text-detection weights.",
    ))
    doc.add_heading("Operational notes", level=2)
    for text in (
        "The classifier and detector checkpoints are committed and load immediately after dependency installation.",
        "EasyOCR weights are downloaded on first OCR use and cached for the running process.",
        "If object detection fails, the app uses OpenCV; if OCR fails, the app preserves the model result and displays a warning.",
        "Uploaded images are processed in memory by application code and are not intentionally persisted.",
    ):
        add_list_item(doc, text, bullet_id)


def add_next_work(doc: Document, bullet_id: int) -> None:
    doc.add_heading("11. Recommended next work", level=1)
    doc.add_paragraph(
        "The next meaningful improvement is not another synthetic-only training run. It is a grouped real-photo dataset that represents the actual deployment domain."
    )
    add_table(
        doc,
        ["Priority", "Action", "Acceptance evidence"],
        [
            ["1", "Collect 50 physical receipts in all four orientations and two scenes", "About 400 labeled real photos"],
            ["2", "Split by physical receipt identity", "No group overlap across train/val/test"],
            ["3", "Fine-tune the prepared 384-pixel MobileNet path", "Improved per-class real-photo recall"],
            ["4", "Recalibrate override and abstention thresholds", "Locked policy before untouched test"],
            ["5", "Report latency and abstention coverage", "p50/p95 time, OCR coverage, uncertainty rate"],
        ],
        [1080, 4680, 3600],
    )
    doc.add_heading("Collection checklist", level=2)
    for text in (
        "Thai and English receipts from different merchants and printers.",
        "Multiple phones, distances, and camera resolutions.",
        "Hands, folds, glare, shadows, blur, and perspective distortion.",
        "Plain and patterned backgrounds.",
        "One stable receipt_group identifier per physical receipt.",
        "A final external holdout kept untouched until all thresholds are frozen.",
    ):
        add_list_item(doc, text, bullet_id)


def add_appendix(doc: Document, bullet_id: int) -> None:
    doc.add_page_break()
    doc.add_heading("Appendix A. Key files and artifacts", level=1)
    for text in (
        "README.md - repository overview, results, quick start, and documentation links.",
        "MODEL_CARD.md - intended use, performance, risks, and mitigations.",
        "docs/METHODOLOGY.md - detailed problem, data, model, and OCR method.",
        "docs/REPRODUCIBILITY.md - complete training and evaluation sequence.",
        "docs/DEPLOYMENT.md - Streamlit Community Cloud build and troubleshooting.",
        "reports/training - aggregate model comparisons, histories, curves, and confusion matrices.",
        "reports/hybrid - aggregate threshold search, summaries, and hybrid confusion matrices.",
    ):
        add_list_item(doc, text, bullet_id)

    doc.add_heading("Appendix B. References", level=1)
    references = (
        "Howard, A. et al. Searching for MobileNetV3. ICCV 2019.",
        "PyTorch and Torchvision documentation: https://pytorch.org/docs/stable/",
        "OpenCV documentation: https://docs.opencv.org/",
        "EasyOCR project: https://github.com/JaidedAI/EasyOCR",
        "Ultralytics documentation: https://docs.ultralytics.com/",
        "Streamlit Community Cloud documentation: https://docs.streamlit.io/deploy/streamlit-community-cloud",
    )
    for reference in references:
        add_list_item(doc, reference, bullet_id)

    add_callout(
        doc,
        "Final note",
        "This report documents an educational computer vision system. It does not validate receipt authenticity, financial values, payment status, or tax compliance.",
        RED,
    )


def build_report() -> None:
    doc = Document()
    for section in doc.sections:
        configure_page(section)
        configure_header_footer(section)
    configure_styles(doc)
    bullet_id = add_numbering_definition(doc, "bullet")

    properties = doc.core_properties
    properties.title = "Receipt Orientation Classifier Project Report"
    properties.subject = "Computer vision project implementation and deployment guide"
    properties.author = "Aung Sann Thit"
    properties.keywords = "computer vision, receipt orientation, MobileNetV3, OCR, Streamlit"

    add_cover(doc)
    add_section_roadmap(doc, bullet_id)
    doc.add_page_break()
    add_executive_summary(doc, bullet_id)
    add_problem_and_labels(doc)
    add_data_pipeline(doc, bullet_id)
    add_preprocessing(doc, bullet_id)
    add_training(doc, bullet_id)
    add_ocr(doc, bullet_id)
    add_app(doc)
    add_code_guide(doc, bullet_id)
    add_results(doc, bullet_id)
    add_testing_and_deployment(doc, bullet_id)
    add_next_work(doc, bullet_id)
    add_appendix(doc, bullet_id)

    for section in doc.sections:
        configure_page(section)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_report()
